#!/usr/bin/env python3
import os
import sys
import json
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# --- CORES INSTITUCIONAIS DA TECHTINS EMPRESA JÚNIOR ---
COLOR_PRIMARY = RGBColor(75, 43, 126)     # Roxo (#4B2B7E)
COLOR_SECONDARY = RGBColor(254, 217, 155) # Amarelo (#FED99B)
COLOR_TEXT = RGBColor(44, 44, 44)         # Preto Fosco (#2C2C2C)
COLOR_GREY = RGBColor(120, 120, 120)     # Cinza Neutro (#787878)

HEX_PRIMARY = "4B2B7E"
HEX_SECONDARY = "FED99B"
HEX_GREY = "CCCCCC"

def set_cell_background(cell, hex_color):
    """Define a cor de fundo de uma célula da tabela."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Define o preenchimento (padding) interno de uma célula da tabela em DXA."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, hex_color=HEX_GREY):
    """Aplica bordas horizontais limpas e remove as verticais para visual moderno."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{hex_color}"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def add_page_number_field(run, field_name):
    """Injeta código XML do Word para campos dinâmicos (PAGE e NUMPAGES)."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_name
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def configurar_estilos(doc):
    """Configura os estilos globais do documento utilizando Times New Roman e cores oficiais."""
    styles = doc.styles
    
    # 1. Estilo Normal (Texto do Corpo)
    style_normal = styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)
    
    # 2. Heading 1 (Título Principal)
    if 'Heading 1' in styles:
        style_h1 = styles['Heading 1']
    else:
        style_h1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(20)
    style_h1.font.bold = True
    style_h1.font.color.rgb = COLOR_PRIMARY
    style_h1.paragraph_format.space_before = Pt(0)
    style_h1.paragraph_format.space_after = Pt(12)
    style_h1.paragraph_format.keep_with_next = True
    
    # 3. Heading 2 (Títulos das Seções)
    if 'Heading 2' in styles:
        style_h2 = styles['Heading 2']
    else:
        style_h2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(13)
    style_h2.font.bold = True
    style_h2.font.color.rgb = COLOR_PRIMARY
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(6)
    style_h2.paragraph_format.keep_with_next = True

def formatar_data(data_iso):
    """Converte data AAAA-MM-DD para DD/MM/AAAA."""
    try:
        partes = data_iso.split("-")
        if len(partes) == 3:
            return f"{partes[2]}/{partes[1]}/{partes[0]}"
    except Exception:
        pass
    return data_iso

def substituir_placeholders(doc, dados):
    """Substitui placeholders padrão no corpo, cabeçalhos, rodapés e tabelas do documento."""
    placeholders = {
        "{{TITULO}}": dados.get("titulo", "ATA DE REUNIÃO"),
        "{{DATA}}": formatar_data(dados.get("data", "")),
        "{{LOCAL}}": dados.get("local", ""),
        "{{PARTICIPANTES}}": ", ".join(dados.get("participantes", []))
    }
    
    # Substitui no corpo
    for p in doc.paragraphs:
        for key, val in placeholders.items():
            if key in p.text:
                if key == "{{PARTICIPANTES}}":
                    p.paragraph_format.space_after = Pt(16)
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                        
    # Substitui em tabelas existentes
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in placeholders.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
                                    
    # Substitui em cabeçalhos e rodapés de todas as seções
    for section in doc.sections:
        for h_f in [section.header, section.footer]:
            if h_f:
                for p in h_f.paragraphs:
                    for key, val in placeholders.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, val)

def aplicar_times_new_roman_global(doc):
    """Garante que todo o texto do documento use Times New Roman."""
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
    for section in doc.sections:
        for h_f in [section.header, section.footer]:
            if h_f:
                for p in h_f.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'

def add_paragraph_with_style_fallback(doc, text="", style_name="Normal", bullet=False):
    """Adiciona um parágrafo com suporte a fallback caso o estilo não exista no documento."""
    try:
        p = doc.add_paragraph(style=style_name)
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
        return p
    except KeyError:
        p = doc.add_paragraph()
        if bullet:
            run_b = p.add_run("• ")
            run_b.font.name = 'Times New Roman'
        if text:
            if style_name.startswith("Heading"):
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Times New Roman'
                if style_name == "Heading 1":
                    run.font.size = Pt(18)
                elif style_name == "Heading 2":
                    run.font.size = Pt(13)
            else:
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
        return p

def template_possui_placeholder(doc, placeholder="{{TITULO}}"):
    """Verifica se o template contém um determinado placeholder no texto."""
    for p in doc.paragraphs:
        if placeholder in p.text:
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        return True
    return False

def main():
    parser = argparse.ArgumentParser(
        description="Gera uma ata formal em Word (.docx) a partir de um arquivo de dados JSON."
    )
    parser.add_argument("-i", "--input", required=True, help="Caminho para o arquivo de ata estruturada JSON.")
    parser.add_argument("-o", "--output", help="Caminho do arquivo .docx final (padrão: mesma pasta da ata de entrada).")
    parser.add_argument("-t", "--template", default="templates/template.docx", help="Caminho para o template .docx base.")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Erro: Arquivo JSON de entrada '{args.input}' não encontrado.")
        sys.exit(1)
        
    with open(args.input, "r", encoding="utf-8") as f:
        dados = json.load(f)
        
    # Inicializa o documento docx
    usa_template = False
    has_placeholders = False
    
    if os.path.exists(args.template):
        doc = Document(args.template)
        usa_template = True
        has_placeholders = template_possui_placeholder(doc, "{{TITULO}}")
        print(f"Utilizando template base: '{args.template}' (Contém placeholders: {has_placeholders})")
        substituir_placeholders(doc, dados)
    else:
        doc = Document()
        print("Nenhum template encontrado. Gerando documento do zero...")
        
    # Se não estiver usando um template pronto, configuramos o layout A4 padrão e estilos
    if not usa_template:
        section = doc.sections[0]
        section.page_width = Inches(8.27)  # A4 Width
        section.page_height = Inches(11.69) # A4 Height
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        configurar_estilos(doc)
        
    # Se o template NÃO possui placeholders declarados (é apenas um layout com cabeçalho/rodapé vazio),
    # ou se estamos criando um documento do zero, geramos o cabeçalho e título manuais
    if not has_placeholders:
        if not usa_template:
            # --- CABEÇALHO DA PRIMEIRA PÁGINA (ORGANIZAÇÃO) ---
            p_header = doc.add_paragraph()
            p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_header_run = p_header.add_run("EMPRESA JÚNIOR DE DESENVOLVIMENTO DE SOFTWARE")
            p_header_run.font.size = Pt(8.5)
            p_header_run.font.bold = True
            p_header_run.font.color.rgb = COLOR_GREY
            
        # --- TÍTULO DO DOCUMENTO ---
        p_title = add_paragraph_with_style_fallback(doc, style_name='Heading 1')
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.runs[0] if p_title.runs else p_title.add_run()
        run_title.text = dados.get("titulo", "ATA DE REUNIÃO").upper()
        run_title.bold = True
        
        # Linha divisória fina sob o título
        p_border = doc.add_paragraph()
        p_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_border.paragraph_format.space_before = Pt(0)
        p_border.paragraph_format.space_after = Pt(16)
        p_border_run = p_border.add_run("━" * 50)
        p_border_run.font.size = Pt(8)
        p_border_run.font.color.rgb = COLOR_PRIMARY
        
        # --- METADADOS DA REUNIÃO (DATA / LOCAL) ---
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(14)
        run_date_label = p_meta.add_run("Data da Reunião: ")
        run_date_label.bold = True
        p_meta.add_run(f"{formatar_data(dados.get('data', ''))}\n")
        
        run_loc_label = p_meta.add_run("Local / Canal: ")
        run_loc_label.bold = True
        p_meta.add_run(f"{dados.get('local', '')}")
        
        # --- PARTICIPANTES ---
        add_paragraph_with_style_fallback(doc, "PARTICIPANTES PRESENTES", style_name='Heading 2')
        participantes_str = ", ".join(dados.get("participantes", []))
        p_part = doc.add_paragraph()
        p_part.add_run(participantes_str)
        p_part.paragraph_format.space_after = Pt(16)
        
    # --- RESUMO DOS ASSUNTOS DISCUTIDOS ---
    add_paragraph_with_style_fallback(doc, "RESUMO DOS ASSUNTOS DISCUTIDOS", style_name='Heading 2')
    
    for item in dados.get("resumo_assuntos", []):
        p_tema = doc.add_paragraph()
        p_tema.paragraph_format.space_before = Pt(6)
        p_tema.paragraph_format.space_after = Pt(2)
        run_tema = p_tema.add_run(f"• {item.get('tema', '')}")
        run_tema.bold = True
        run_tema.font.color.rgb = COLOR_PRIMARY
        
        p_det = doc.add_paragraph()
        p_det.paragraph_format.left_indent = Inches(0.25)
        p_det.add_run(item.get("detalhes", ""))
        
    # --- TABELA DE ENCAMINHAMENTOS ---
    add_paragraph_with_style_fallback(doc, "TABELA DE ENCAMINHAMENTOS (AÇÕES E PRAZOS)", style_name='Heading 2')
    
    encaminhamentos = dados.get("encaminhamentos", [])
    if encaminhamentos:
        # Cria tabela com 3 colunas (Ação, Responsável, Prazo)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Largura das colunas (Soma = ~6.27 polegadas de área útil na página A4)
        col_widths = [Inches(3.27), Inches(1.5), Inches(1.5)]
        
        # Cabeçalhos da Tabela
        hdr_cells = table.rows[0].cells
        headers_labels = ["Ação / Atividade", "Responsável", "Prazo"]
        
        for i, text in enumerate(headers_labels):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) # Texto Branco
            set_cell_background(hdr_cells[i], HEX_PRIMARY) # Fundo Azul Escuro
            set_cell_margins(hdr_cells[i], top=140, bottom=140, left=140, right=140)
            hdr_cells[i].width = col_widths[i]
            
        # Conteúdo da Tabela
        for idx, enc in enumerate(encaminhamentos):
            row_cells = table.add_row().cells
            valores = [enc.get("acao", ""), enc.get("responsavel", ""), enc.get("prazo", "")]
            
            for i, val in enumerate(valores):
                row_cells[i].text = ""
                p = row_cells[i].paragraphs[0]
                p.add_run(val)
                set_cell_margins(row_cells[i], top=100, bottom=100, left=140, right=140)
                row_cells[i].width = col_widths[i]
                
                # Zebrador simples (linhas pares com fundo roxo claro de destaque)
                if idx % 2 != 0:
                    set_cell_background(row_cells[i], "F1EDF8")
                    
        set_table_borders(table)
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(0)
        p_space.paragraph_format.space_after = Pt(12)
    else:
        p_no_enc = doc.add_paragraph("Nenhum encaminhamento ou pendência registrada para esta reunião.")
        p_no_enc.paragraph_format.space_after = Pt(16)
        
    # --- REGISTRO OBRIGATÓRIO DE DÚVIDAS / INCERTEZAS ---
    add_paragraph_with_style_fallback(doc, "DÚVIDAS, ANOTAÇÕES E INCERTEZAS", style_name='Heading 2')
    
    duvidas = dados.get("duvidas_incertezas", [])
    if duvidas:
        for duv in duvidas:
            p_duv = add_paragraph_with_style_fallback(doc, duv, style_name='List Bullet', bullet=True)
            if p_duv.runs:
                p_duv.runs[-1].italic = True
    else:
        doc.add_paragraph("Nenhuma dúvida ou incerteza registrada.")
        
    # --- STATUS DE DESENVOLVIMENTO (GITHUB) ---
    github_status = dados.get("status_desenvolvimento", [])
    if github_status:
        p_space_gh = doc.add_paragraph()
        p_space_gh.paragraph_format.space_before = Pt(14)
        p_space_gh.paragraph_format.space_after = Pt(0)
        
        add_paragraph_with_style_fallback(doc, "STATUS DE DESENVOLVIMENTO (GITHUB)", style_name='Heading 2')
        
        table_gh = doc.add_table(rows=1, cols=4)
        table_gh.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        col_widths_gh = [Inches(1.5), Inches(1.2), Inches(2.3), Inches(1.2)]
        
        hdr_cells_gh = table_gh.rows[0].cells
        headers_labels_gh = ["Repositório", "Atividade Recente", "Último Commit / Detalhes", "Última Atualização"]
        
        for i, text in enumerate(headers_labels_gh):
            hdr_cells_gh[i].text = ""
            p = hdr_cells_gh[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hdr_cells_gh[i], HEX_PRIMARY)
            set_cell_margins(hdr_cells_gh[i], top=140, bottom=140, left=140, right=140)
            hdr_cells_gh[i].width = col_widths_gh[i]
            
        for idx, item in enumerate(github_status):
            row_cells_gh = table_gh.add_row().cells
            valores_gh = [
                item.get("repositorio", ""),
                str(item.get("commits_recentes", "Sem commits")),
                item.get("ultimo_commit", "Nenhuma informação"),
                item.get("data", "")
            ]
            
            for i, val in enumerate(valores_gh):
                row_cells_gh[i].text = ""
                p = row_cells_gh[i].paragraphs[0]
                p.add_run(val)
                set_cell_margins(row_cells_gh[i], top=100, bottom=100, left=140, right=140)
                row_cells_gh[i].width = col_widths_gh[i]
                
                if idx % 2 != 0:
                    set_cell_background(row_cells_gh[i], "F1EDF8")
                    
        set_table_borders(table_gh)
        p_space_after_gh = doc.add_paragraph()
        p_space_after_gh.paragraph_format.space_before = Pt(0)
        p_space_after_gh.paragraph_format.space_after = Pt(12)
        
    # Se não estiver usando um template pronto, configuramos o rodapé padrão com paginação
    if not usa_template:
        section = doc.sections[0]
        footer = section.footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_footer.text = "" # Limpa qualquer texto padrão
        
        # Nota de Confidencialidade no canto esquerdo, paginação no canto direito (usando tab stop)
        p_footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.27))
        run_conf = p_footer.add_run("Documento de Uso Interno - EJ\tPágina ")
        run_conf.font.size = Pt(8.5)
        run_conf.font.color.rgb = COLOR_GREY
        
        add_page_number_field(p_footer.add_run(), "PAGE")
        p_footer.add_run(" de ").font.color.rgb = COLOR_GREY
        add_page_number_field(p_footer.add_run(), "NUMPAGES")
        
        for r in p_footer.runs:
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_GREY
        
    # --- APLICA FONTE TIMES NEW ROMAN EM TODO O DOCUMENTO ---
    aplicar_times_new_roman_global(doc)
    
    # --- SALVA O ARQUIVO ---
    output_path = args.output
    if not output_path:
        output_path = os.path.splitext(args.input)[0] + ".docx"
        
    doc.save(output_path)
    print(f"Sucesso! Documento Word oficial gerado e salvo em '{output_path}'.")

if __name__ == "__main__":
    main()
